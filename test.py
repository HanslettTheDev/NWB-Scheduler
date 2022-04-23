from docx import Document
from docx.shared import Inches

# document.add_heading('Document Title', 0)

# p = document.add_paragraph('A plain paragraph having some ')
# p.add_run('bold').bold = True
# p.add_run(' and some ')
# p.add_run('italic.').italic = True

# document.add_heading('Heading, level 1', level=1)
# document.add_paragraph('Intense quote', style='Intense Quote')

# document.add_paragraph(
#     'first item in unordered list', style='List Bullet'
# )
# document.add_paragraph(
#     'first item in ordered list', style='List Number'
# )

class CreateTables:
	def __init__(self, row, col, *args):
		self.row = row
		self.col = col
		self.headers = list(args)
		self.f = open("demo.docx", 'ab+')
		self.document = Document(self.f)

	def createRowsAndColumns(self):
		table = self.document.add_table(rows=self.row, cols=self.col)
		hdr_cells = table.rows[0].cells
		print(self.row, self.col, self.headers)
		counter = self.col - 1
		i = 0
		while i <= counter:
			hdr_cells[i].text = self.headers[i]
			i += 1
		self.document.save("demo.docx")


my_tables = CreateTables(1, 3, 'JANUARY 4-10|LEVITICUS 18-19', 'Chairman:', 'Name')
my_tables2 = CreateTables(2, 4, '6:30', 'Song 122:', 'Prayer:', 'Brother')

my_tables.createRowsAndColumns()
my_tables2.createRowsAndColumns()

# my_tables2 = CreateTables(1, 3, 'JANUARY 4-10|LEVITICUS 18-19', 'Chairman:', 'Name')
# my_tables2.createRowsAndColumns()


# records = (
#     (3, '101', 'Spam'),
#     (7, '422', 'Eggs'),
#     (4, '631', 'Spam, spam, eggs, and spam')
# )

# table = document.add_table(rows=1, cols=3)
# hdr_cells = table.rows[0].cells
# hdr_cells[0].text = 'JANUARY 4-10|LEVITICUS 18-19'
# hdr_cells[1].text = 'Chairman:'
# hdr_cells[2].text = 'Name'

# for qty, id, desc in records:
#     row_cells = table.add_row().cells
#     row_cells[0].text = str(qty)
#     row_cells[1].text = id
#     row_cells[2].text = desc

# document.add_page_break()

# document.save('demo.docx')